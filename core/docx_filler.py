"""Заполнение шаблона .docx с подстановкой плейсхолдеров вида {{КЛЮЧ}}.

Word нередко разбивает плейсхолдер на несколько runs (особенно если
ставился точечный bold/курсив или были автоисправления). Поэтому замена
делается на уровне всего параграфа: текст runs объединяется, заменяется,
и пишется обратно в первый run; остальные очищаются — форматирование
первого run сохраняется.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

PLACEHOLDER_RE = re.compile(r"\{\{([A-Za-zА-Яа-я0-9_]+)\}\}")


def _replace_in_paragraph(paragraph: Paragraph, context: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    def repl(match: re.Match[str]) -> str:
        key = match.group(1)
        return str(context.get(key, match.group(0)))

    new_text = PLACEHOLDER_RE.sub(repl, full_text)
    if new_text == full_text:
        return

    # Записываем результат в первый run, остальные очищаем.
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _iter_paragraphs(parent) -> Iterable[Paragraph]:
    """Рекурсивный обход параграфов: сам объект, его таблицы, ячейки таблиц."""
    if hasattr(parent, "paragraphs"):
        yield from parent.paragraphs
    if hasattr(parent, "tables"):
        for table in parent.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_all_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    yield from _iter_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_paragraphs(section.header)
        yield from _iter_paragraphs(section.footer)


def fill_template(template_path: Path | str, context: dict[str, str], output_path: Path | str) -> list[str]:
    """Заполнить шаблон .docx и сохранить результат.

    Возвращает список незаменённых плейсхолдеров (если такие остались).
    Это warnings — ai_validator потом подтвердит/уточнит.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    str_context = {k: str(v) if v is not None else "" for k, v in context.items()}

    for paragraph in _iter_all_paragraphs(doc):
        _replace_in_paragraph(paragraph, str_context)

    doc.save(str(output_path))

    return _find_remaining_placeholders(output_path)


def _find_remaining_placeholders(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    remaining: set[str] = set()
    for paragraph in _iter_all_paragraphs(doc):
        for match in PLACEHOLDER_RE.finditer(paragraph.text):
            remaining.add(match.group(0))
    return sorted(remaining)


def extract_text(docx_path: Path | str) -> str:
    """Извлечь весь текст документа для AI-валидации."""
    doc = Document(str(docx_path))
    parts: list[str] = []
    for paragraph in _iter_all_paragraphs(doc):
        if paragraph.text.strip():
            parts.append(paragraph.text)
    return "\n".join(parts)
