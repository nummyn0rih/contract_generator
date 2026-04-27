"""Сгенерировать тестовые data/farmers.xlsx и data/template.docx.

Запуск: python scripts/make_samples.py

Генерирует:
- 10 фермеров: 7 валидных + 3 со специально внесёнными ошибками
  (битый ИНН, отсутствует БИК, отрицательный объём).
- Шаблон договора с плейсхолдерами, которые поддерживает docx_filler.

Все ИНН/БИК/счёта в валидных строках проходят контрольные суммы
(используем реальный БИК Сбербанка и реальный корсчёт).
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from core import validators as v  # noqa: E402

DATA_DIR = ROOT / "data"
DATA_DIR.mkdir(exist_ok=True)

# Реальный БИК и корсчёт Сбербанка (г. Москва).
SBER_BIK = "044525225"
SBER_CORR = "30101810400000000225"


_INN10_W = (2, 4, 10, 3, 5, 9, 4, 6, 8)
_INN12_W11 = (7, 2, 4, 10, 3, 5, 9, 4, 6, 8)
_INN12_W12 = (3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8)


def make_inn10(prefix9: str) -> str:
    assert len(prefix9) == 9 and prefix9.isdigit()
    digits = [int(c) for c in prefix9]
    csum = sum(d * w for d, w in zip(digits, _INN10_W)) % 11 % 10
    inn = prefix9 + str(csum)
    assert v.validate_inn(inn)
    return inn


def make_inn12(prefix10: str) -> str:
    assert len(prefix10) == 10 and prefix10.isdigit()
    digits = [int(c) for c in prefix10]
    c11 = sum(d * w for d, w in zip(digits, _INN12_W11)) % 11 % 10
    digits.append(c11)
    c12 = sum(d * w for d, w in zip(digits, _INN12_W12)) % 11 % 10
    inn = prefix10 + str(c11) + str(c12)
    assert v.validate_inn(inn)
    return inn


def make_account(bik: str, prefix19: str) -> str:
    """Подобрать 20-ю цифру так, чтобы счёт прошёл контрольную сумму с БИК."""
    assert len(prefix19) == 19 and prefix19.isdigit()
    for last in range(10):
        candidate = prefix19 + str(last)
        if v.validate_account_number(candidate, bik):
            return candidate
    raise RuntimeError(f"не удалось подобрать счёт для prefix={prefix19}")


def build_farmers() -> pd.DataFrame:
    rows: list[dict] = []

    # 7 валидных фермеров.
    valid_inputs = [
        ("Иванов Иван Иванович",        "770100100", "770101001", "Краснодарский край",  "Пшеница",      "150.00",  "12500.00", "2025-05-15"),
        ("ООО Колос",                   "770700100", "770701001", "Ростовская область",  "Подсолнечник", "200.00",  "28000.00", "2025-05-20"),
        ("ИП Петров Сергей Александрович","770200100","",          "Воронежская область", "Кукуруза",     "180.50",  "11000.00", "2025-05-22"),
        ("ООО АгроЮг",                  "773300100", "773301001", "Ставропольский край", "Ячмень",       "320.00",   "9800.00", "2025-06-01"),
        ("КФХ Сидорова",                "770400100", "",          "Белгородская область","Соя",          "95.00",   "32000.00", "2025-06-05"),
        ("ООО Заря",                    "770900100", "770901001", "Тамбовская область",  "Сахарная свёкла","450.00",  "5500.00", "2025-06-10"),
        ("Кузнецов Дмитрий Олегович",   "770500100", "",          "Курская область",     "Рапс",         "75.50",   "27500.00", "2025-06-15"),
    ]

    account_serial = 1
    for full_name, prefix9, kpp, region, culture, volume, price, ddate in valid_inputs:
        inn = (
            make_inn10(prefix9) if kpp else make_inn12(prefix9 + str(account_serial % 10))
        )
        account = make_account(SBER_BIK, f"40702810{account_serial:011d}")
        account_serial += 1
        rows.append(
            {
                "ФИО": full_name,
                "ИНН": inn,
                "КПП": kpp,
                "Банк": "ПАО Сбербанк",
                "БИК": SBER_BIK,
                "Р/счёт": account,
                "К/счёт": SBER_CORR,
                "Регион": region,
                "Культура": culture,
                "Объём, т": volume,
                "Цена за тонну, руб": price,
                "Дата договора": ddate,
                "Сезон": "2025",
                "Особые условия": "",
            }
        )

    # 3 строки с намеренными ошибками — должны попасть в data_error.
    rows.append(
        {
            "ФИО": "Битый ИНН Иваныч",
            "ИНН": "1234567890",  # 10 цифр, но контрольная сумма не сходится
            "КПП": "770101001",
            "Банк": "ПАО Сбербанк",
            "БИК": SBER_BIK,
            "Р/счёт": make_account(SBER_BIK, f"40702810{99:011d}"),
            "К/счёт": SBER_CORR,
            "Регион": "Москва",
            "Культура": "Пшеница",
            "Объём, т": "100.00",
            "Цена за тонну, руб": "10000.00",
            "Дата договора": "2025-05-15",
            "Сезон": "2025",
            "Особые условия": "",
        }
    )
    rows.append(
        {
            "ФИО": "ООО Без БИК",
            "ИНН": make_inn10("770800100"),
            "КПП": "770801001",
            "Банк": "ПАО Сбербанк",
            "БИК": "",  # обязательное поле пустое
            "Р/счёт": "40702810000000000123",
            "К/счёт": SBER_CORR,
            "Регион": "Москва",
            "Культура": "Кукуруза",
            "Объём, т": "100.00",
            "Цена за тонну, руб": "10000.00",
            "Дата договора": "2025-05-15",
            "Сезон": "2025",
            "Особые условия": "",
        }
    )
    rows.append(
        {
            "ФИО": "КФХ Минусовое",
            "ИНН": make_inn10("770600100"),
            "КПП": "",
            "Банк": "ПАО Сбербанк",
            "БИК": SBER_BIK,
            "Р/счёт": make_account(SBER_BIK, f"40702810{100:011d}"),
            "К/счёт": SBER_CORR,
            "Регион": "Тверская область",
            "Культура": "Овёс",
            "Объём, т": "-50.00",  # отрицательный объём
            "Цена за тонну, руб": "8000.00",
            "Дата договора": "2025-05-15",
            "Сезон": "2025",
            "Особые условия": "",
        }
    )

    return pd.DataFrame(rows)


def write_farmers_xlsx(df: pd.DataFrame, path: Path) -> None:
    df.to_excel(path, index=False)
    print(f"Записано: {path} ({len(df)} строк)")


def write_template_docx(path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    run = title.add_run("ДОГОВОР ПОСТАВКИ № ___")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = 1  # center

    doc.add_paragraph("г. Москва\t\t\t\t\t\t\t\t{{ДАТА}}")
    doc.add_paragraph()
    doc.add_paragraph(
        "Покупатель ООО «Заготовитель», именуемое в дальнейшем «Покупатель», "
        "и {{ФИО}} (ИНН {{ИНН}}, КПП {{КПП}}), именуемое в дальнейшем «Поставщик», "
        "заключили настоящий договор о нижеследующем:"
    )

    doc.add_paragraph()
    h = doc.add_paragraph().add_run("1. ПРЕДМЕТ ДОГОВОРА")
    h.bold = True
    doc.add_paragraph(
        "1.1. Поставщик обязуется поставить Покупателю сельскохозяйственную продукцию — "
        "{{КУЛЬТУРА}} (далее — Товар) в объёме {{ОБЪЕМ}} тонн, выращенную в "
        "{{РЕГИОН}} в сезон {{СЕЗОН}}."
    )
    doc.add_paragraph(
        "1.2. Покупатель обязуется принять и оплатить Товар по цене {{ЦЕНА_ЗА_ТОННУ}} руб. за тонну."
    )

    doc.add_paragraph()
    h = doc.add_paragraph().add_run("2. ЦЕНА И ПОРЯДОК РАСЧЁТОВ")
    h.bold = True
    doc.add_paragraph(
        "2.1. Общая сумма договора составляет {{СУММА}} руб. ({{СУММА_ПРОПИСЬЮ}})."
    )
    doc.add_paragraph(
        "2.2. Оплата производится безналичным перечислением на расчётный счёт Поставщика "
        "в течение 10 (десяти) банковских дней с даты приёмки Товара."
    )

    doc.add_paragraph()
    h = doc.add_paragraph().add_run("3. ИНДИВИДУАЛЬНЫЕ УСЛОВИЯ")
    h.bold = True
    doc.add_paragraph("{{ИНДИВИДУАЛЬНЫЕ_УСЛОВИЯ}}")

    doc.add_paragraph()
    h = doc.add_paragraph().add_run("4. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    h.bold = True

    doc.add_paragraph("Поставщик:")
    doc.add_paragraph("{{ФИО}}")
    doc.add_paragraph("ИНН: {{ИНН}}")
    doc.add_paragraph("КПП: {{КПП}}")
    doc.add_paragraph("Банк: {{БАНК}}")
    doc.add_paragraph("БИК: {{БИК}}")
    doc.add_paragraph("Р/счёт: {{Р_СЧЕТ}}")
    doc.add_paragraph("Корсчёт: {{К_СЧЕТ}}")
    doc.add_paragraph()
    doc.add_paragraph("Подпись: ______________________ / {{ФИО}} /")

    doc.save(str(path))
    print(f"Записано: {path}")


def main() -> None:
    df = build_farmers()
    write_farmers_xlsx(df, DATA_DIR / "farmers.xlsx")
    write_template_docx(DATA_DIR / "template.docx")
    print("Готово.")


if __name__ == "__main__":
    main()
